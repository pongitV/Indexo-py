from pathlib import Path
from loguru import logger

def extract_docx_text(docx_path: Path) -> str:
    try:
        import docx
        doc = docx.Document(str(docx_path))
        text_parts = []
        for p in doc.paragraphs:
            if p.text:
                text_parts.append(p.text)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text:
                        text_parts.append(cell.text)
        return "\n".join(text_parts)
    except Exception as e:
        logger.warning("Failed to extract DOCX text from {}: {}", docx_path, e)
        return ""

def extract_odt_text(odt_path: Path) -> str:
    try:
        from odf import opendocument, text, teletype
        doc = opendocument.load(str(odt_path))
        paragraphs = doc.getElementsByType(text.P)
        text_parts = [teletype.extractText(p) for p in paragraphs if teletype.extractText(p)]
        return "\n".join(text_parts)
    except Exception as e:
        logger.warning("Failed to extract ODT text from {}: {}", odt_path, e)
        return ""

def extract_document_text(doc_path: Path) -> str:
    ext = doc_path.suffix.lower()
    if ext == ".docx":
        return extract_docx_text(doc_path)
    elif ext == ".odt":
        return extract_odt_text(doc_path)
    return ""
